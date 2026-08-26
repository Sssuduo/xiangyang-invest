"""工作日历 API 路由"""
import io
import json
import os
from collections import defaultdict
from datetime import datetime, timedelta, timezone

from flask import Blueprint, request, jsonify, send_file, g, current_app

from models.investment import WorkCalendarEntry
from extensions import db
from routes.business_auth import business_login_required

# docx 库用于生成 Word 文档（可选依赖）
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

work_calendar_bp = Blueprint('work_calendar', __name__, url_prefix='/api/work-calendar')
bp = work_calendar_bp  # 别名，方便内部使用

# 业务时区：中国标准时间（无夏令时；存储层统一 UTC，仅在展示/导出时转换）
LOCAL_TZ = timezone(timedelta(hours=8))
# 附件同源前缀：导出只允许读取本站上传目录内的文件（同时防 SSRF）
UPLOAD_PREFIX = '/static/uploads/'
# 单张图片嵌入 Word 的大小上限（10MB），防内存暴涨
MAX_IMAGE_BYTES = 10 * 1024 * 1024
# 列表查询单次返回上限
PAGE_LIMIT = 500
# 合法工作时段
TIME_PERIODS = {'morning', 'afternoon', 'custom'}


# ===== 工具函数 =====

def _to_local(dt):
    """存储的 UTC 时间 → 东八区本地时间；naive 视为 UTC（SQLite 读出均为 naive）。"""
    if dt is None:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(LOCAL_TZ)


def _parse_iso(value):
    """解析 ISO 时间（'Z' 结尾按 UTC），失败抛 ValueError。"""
    return datetime.fromisoformat(value.replace('Z', '+00:00'))


def _collect_entry_data(data):
    """提取并校验创建/更新条目的公共字段；非法抛 ValueError。"""
    work_item = (data.get('work_item') or '').strip()
    if not work_item:
        raise ValueError('工作事项不能为空')
    if len(work_item) > 200:
        raise ValueError('工作事项不能超过 200 字')

    try:
        start_dt = _parse_iso(data['start_datetime'])
        end_dt = _parse_iso(data['end_datetime'])
    except (KeyError, ValueError) as e:
        raise ValueError(f'时间格式错误: {e}')

    if end_dt <= start_dt:
        raise ValueError('结束时间必须晚于开始时间')

    time_period = data.get('time_period')
    if time_period is not None and time_period not in TIME_PERIODS:
        raise ValueError(f'time_period 非法，仅支持 {", ".join(sorted(TIME_PERIODS))}')

    participants = data.get('participants', [])
    attachments = data.get('attachments', [])
    if not isinstance(participants, list):
        raise ValueError('participants 必须是数组')
    if not isinstance(attachments, list):
        raise ValueError('attachments 必须是数组')

    return {
        'start_datetime': start_dt,
        'end_datetime': end_dt,
        'time_period': time_period,
        'work_item': work_item,
        'work_content': (data.get('work_content') or '').strip(),
        'participants': json.dumps(participants, ensure_ascii=False),
        'attachments': json.dumps(attachments, ensure_ascii=False),
    }


# ===== 日历 CRUD =====

@bp.route('', methods=['GET'])
@business_login_required
def get_list():
    """获取工作日历条目（按时间范围，重叠语义：事件与该区间有交集即返回）"""
    start = request.args.get('start')
    end = request.args.get('end')
    user_id = g.user.id

    query = WorkCalendarEntry.query.filter_by(user_id=user_id)

    if start and end:
        try:
            start_dt = _parse_iso(start)
            end_dt = _parse_iso(end)
        except ValueError:
            return jsonify({'code': 400, 'message': '日期格式错误'}), 400
        # 区间重叠：事件在 [start_dt, end_dt] 内任意一段相交即返回，避免跨边界事件丢失
        query = query.filter(
            WorkCalendarEntry.start_datetime <= end_dt,
            WorkCalendarEntry.end_datetime >= start_dt
        )
    elif start or end:
        return jsonify({'code': 400, 'message': 'start 与 end 必须同时提供'}), 400

    entries = query.order_by(WorkCalendarEntry.start_datetime).limit(PAGE_LIMIT).all()
    return jsonify({'code': 0, 'data': [e.to_dict() for e in entries]})


@bp.route('', methods=['POST'])
@business_login_required
def create():
    """创建工作日历条目"""
    data = request.json or {}

    try:
        fields = _collect_entry_data(data)
    except ValueError as e:
        return jsonify({'code': 400, 'message': str(e)}), 400

    entry = WorkCalendarEntry(
        user_id=g.user.id,
        start_datetime=fields['start_datetime'],
        end_datetime=fields['end_datetime'],
        time_period=fields['time_period'],
        work_item=fields['work_item'],
        work_content=fields['work_content'],
        participants=fields['participants'],
        attachments=fields['attachments'],
        created_by=g.user.id
    )

    db.session.add(entry)
    db.session.commit()

    return jsonify({'code': 0, 'data': entry.to_dict()}), 201


@bp.route('/<int:id>', methods=['PUT'])
@business_login_required
def update(id):
    """更新工作日历条目（仅限本人记录）"""
    entry = WorkCalendarEntry.query.filter_by(id=id, user_id=g.user.id).first()
    if entry is None:
        return jsonify({'code': 404, 'message': '记录不存在'}), 404

    data = request.json or {}

    try:
        fields = _collect_entry_data(data)
    except ValueError as e:
        return jsonify({'code': 400, 'message': str(e)}), 400

    entry.start_datetime = fields['start_datetime']
    entry.end_datetime = fields['end_datetime']
    entry.time_period = fields['time_period']
    entry.work_item = fields['work_item']
    entry.work_content = fields['work_content']
    entry.participants = fields['participants']
    entry.attachments = fields['attachments']
    entry.updated_at = datetime.utcnow()

    db.session.commit()

    return jsonify({'code': 0, 'data': entry.to_dict()})


@bp.route('/<int:id>', methods=['DELETE'])
@business_login_required
def delete(id):
    """删除工作日历条目（仅限本人记录）"""
    entry = WorkCalendarEntry.query.filter_by(id=id, user_id=g.user.id).first()
    if entry is None:
        return jsonify({'code': 404, 'message': '记录不存在'}), 404

    db.session.delete(entry)
    db.session.commit()

    return jsonify({'code': 0, 'message': '删除成功'})


# ===== Word 导出 =====

def _fmt(run, size=11, bold=False, color=None):
    """统一中文字体设置"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color is not None:
        run.font.color.rgb = color
    return run


def _para(doc, text='', size=11, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        _fmt(p.add_run(text), size=size, bold=bold, color=color)
    return p


def _resolve_local_image(url):
    """把同源附件 URL 解析为本地文件路径；非本站文件/非法路径返回 None（防 SSRF）。"""
    if not url or not url.startswith(UPLOAD_PREFIX):
        return None
    rel = url[len(UPLOAD_PREFIX):]
    if not rel or '..' in rel.replace('\\', '/'):
        return None
    upload_dir = os.path.normpath(current_app.config['UPLOAD_FOLDER'])
    path = os.path.normpath(os.path.join(upload_dir, rel))
    if not path.startswith(upload_dir):
        return None
    return path


def _is_image(url):
    if not url:
        return False
    ext = os.path.splitext(url.split('?')[0])[1].lower()
    return ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']


@bp.route('/export/word', methods=['POST'])
@business_login_required
def export_word():
    """导出工作日历为 Word 文档"""
    if not DOCX_AVAILABLE:
        return jsonify({'code': 500, 'message': 'Word 导出功能不可用，请安装 docx 库'}), 500

    data = request.json or {}
    start = data.get('start')
    end = data.get('end')

    if not start or not end:
        return jsonify({'code': 400, 'message': '请提供开始和结束日期'}), 400

    try:
        start_dt = _parse_iso(start)
        end_dt = _parse_iso(end)
    except ValueError:
        return jsonify({'code': 400, 'message': '日期格式错误'}), 400

    # 导出字段过滤：fields 可选（time/work_content/participants/attachments），缺省/空 = 全输出
    EXPORT_FIELDS = {'time', 'work_content', 'participants', 'attachments'}
    fields = set(data.get('fields') or []) & EXPORT_FIELDS
    if not fields:
        fields = set(EXPORT_FIELDS)

    # 重叠语义查询该时间段内有交集的条目
    entries = WorkCalendarEntry.query.filter(
        WorkCalendarEntry.user_id == g.user.id,
        WorkCalendarEntry.start_datetime <= end_dt,
        WorkCalendarEntry.end_datetime >= start_dt
    ).order_by(WorkCalendarEntry.start_datetime).all()

    if not entries:
        return jsonify({'code': 404, 'message': '该时间段内没有工作记录'}), 404

    # 生成 Word 文档
    doc = Document()

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题（用查询范围边界，避免与循环内变量混淆）
    _para(doc, '工作日历记录', size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f'（{_to_local(start_dt).strftime("%Y-%m-%d")} 至 {_to_local(end_dt).strftime("%Y-%m-%d")}）',
          size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()  # 空行

    # 按日期分组（按东八区本地日期归类）
    entries_by_date = defaultdict(list)
    for entry in entries:
        date_key = _to_local(entry.start_datetime).strftime('%Y-%m-%d')
        entries_by_date[date_key].append(entry)

    # 星期映射
    weekday_map = {
        0: '周一', 1: '周二', 2: '周三', 3: '周四',
        4: '周五', 5: '周六', 6: '周日'
    }

    # 遍历每个日期
    for date_str in sorted(entries_by_date.keys()):
        date_entries = entries_by_date[date_str]
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        weekday = weekday_map[date_obj.weekday()]

        # 日期标题
        _para(doc, f'【{date_str} {weekday}】', size=16, bold=True)

        # 遍历该日期的每条记录
        for idx, entry in enumerate(date_entries):
            _para(doc, f'{idx + 1}. {entry.work_item}', size=14, bold=True)

            # 时间（东八区）
            if 'time' in fields:
                start_local = _to_local(entry.start_datetime)
                end_local = _to_local(entry.end_datetime)
                time_str = f"{start_local.strftime('%H:%M')} - {end_local.strftime('%H:%M')}"
                period_label = {'morning': '(上午)', 'afternoon': '(下午)'}.get(entry.time_period, '')
                time_para = doc.add_paragraph()
                _fmt(time_para.add_run('工作时段：'), size=11, bold=True)
                _fmt(time_para.add_run(f"{time_str} {period_label}"), size=11)

            # 工作内容
            if 'work_content' in fields:
                content_para = doc.add_paragraph()
                _fmt(content_para.add_run('工作内容：'), size=11, bold=True)
                _fmt(content_para.add_run(entry.work_content or '无'), size=11)

            # 参加人员
            if 'participants' in fields:
                participants = json.loads(entry.participants) if entry.participants else []
                participants_para = doc.add_paragraph()
                _fmt(participants_para.add_run('参加人员：'), size=11, bold=True)
                _fmt(participants_para.add_run('、'.join(participants) if participants else '无'), size=11)

            # 附件（仅嵌入本站图片附件，外部链接只列文件名不下载）
            if 'attachments' in fields:
                attachments = json.loads(entry.attachments) if entry.attachments else []
                image_attachments = [att for att in attachments if _is_image(att.get('url', ''))]

                attachments_para = doc.add_paragraph()
                _fmt(attachments_para.add_run('附件：'), size=11, bold=True)

                if image_attachments:
                    _fmt(attachments_para.add_run(f'共 {len(image_attachments)} 张图片（已嵌入文档）'), size=11)

                    # 嵌入每张图片（仅读取本站 uploads 目录内文件）
                    for att in image_attachments:
                        url = att.get('url', '')
                        name = att.get('name', '图片')
                        local_path = _resolve_local_image(url)
                        try:
                            if not local_path or not os.path.isfile(local_path):
                                raise OSError('附件不在本站上传目录或文件不存在')
                            if os.path.getsize(local_path) > MAX_IMAGE_BYTES:
                                raise OSError('图片超过 10MB，已跳过')
                            with open(local_path, 'rb') as f:
                                image_data = f.read()
                            if not image_data:
                                raise OSError('图片内容为空')

                            doc.add_picture(io.BytesIO(image_data), width=Inches(4))

                            caption_para = doc.add_paragraph()
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _fmt(caption_para.add_run(f'图：{name}'), size=9,
                                 color=RGBColor(0x80, 0x80, 0x80))
                        except Exception as e:
                            current_app.logger.warning('加载图片失败: %s, 错误: %s', url, e)
                            error_para = doc.add_paragraph()
                            error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _fmt(error_para.add_run(f'[图片加载失败: {name}]'), size=9,
                                 color=RGBColor(0x80, 0x80, 0x80))
                elif attachments:
                    # 有附件但没有本站图片
                    non_image_names = [att.get('name', '附件') for att in attachments
                                       if not _is_image(att.get('url', ''))]
                    _fmt(attachments_para.add_run(f'非图片附件：{", ".join(non_image_names)}（未嵌入）'), size=11)
                else:
                    _fmt(attachments_para.add_run('无'), size=11)

            # 记录之间的分隔线
            doc.add_paragraph('─' * 40)

        doc.add_paragraph()  # 日期之间的空行

    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"工作日历_{_to_local(start_dt).strftime('%Y%m%d')}_至_{_to_local(end_dt).strftime('%Y%m%d')}.docx"

    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )