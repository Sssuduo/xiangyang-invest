"""
公文 .docx 生成服务

将模型返回（或用户提供的）Markdown 风格文本转换为 .docx，
保存在 instance/generated_docs/ 下，返回文件名供下载路由签发。
依赖 python-docx（项目读取 .docx 时同款依赖，服务器已安装）。
"""
import os
import re
import uuid

from flask import current_app

_GEN_DIRNAME = 'generated_docs'


def _ensure_dir():
    base = os.path.join(current_app.instance_path, _GEN_DIRNAME)
    os.makedirs(base, exist_ok=True)
    return base


def _safe_filename(title):
    # 去掉文件名非法字符，限制长度
    name = re.sub(r'[\\/:*?"<>|\r\n]+', '_', title or '公文文档').strip()
    name = name[:60] or '公文文档'
    return f"{name}_{uuid.uuid4().hex[:8]}.docx"


def generate_docx(content, title='公文文档'):
    """生成 .docx，返回文件名（basename）"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    heading = doc.add_heading(level=0)
    run = heading.add_run(title or '公文文档')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for raw in (content or '').split('\n'):
        line = raw.rstrip()
        if not line.strip():
            continue

        m = re.match(r'^(#{1,3})\s+(.*)$', line)
        if m:
            level = len(m.group(1))  # 1/2/3
            h = doc.add_heading(level=level)
            h.add_run(m.group(2).strip())
            continue

        # 普通段落：首行缩进 2 字符
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.add_run(line.strip())

    out_dir = _ensure_dir()
    filename = _safe_filename(title)
    out_path = os.path.join(out_dir, filename)
    doc.save(out_path)
    return filename
