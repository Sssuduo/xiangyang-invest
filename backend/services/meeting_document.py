"""
活动台账录音 Word 文档生成

为已识别的活动台账生成包含「分段原文 + 清洁版 + 摘要版」三部分的 Word 文档。
复用 backend/services/word_service.py 的 markdown → docx 解析风格。
"""
import os
import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _ensure_dir():
    """确保输出目录存在，返回 static/meetings 目录绝对路径"""
    base = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'meetings')
    os.makedirs(base, exist_ok=True)
    return base


def _strip_markdown(text):
    """移除 markdown 标记，保留纯文本"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def generate_meeting_docx(activity, segmented_text='', clean_text='', summary_text=''):
    """为活动台账生成三件套 docx: 发言分段 + 清洁版 + 摘要版。

    Args:
        activity: ActivityLedger 实例
        segmented_text: 分段后原文
        clean_text: 清洁版文本
        summary_text: 摘要版文本

    Returns:
        tuple: (static_url 相对路径如 '/static/meetings/xxx.docx', file_name)
    """
    out_dir = _ensure_dir()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', (activity.content or '活动台账')[:30])
    file_name = f'{safe_name}_会议总结_{timestamp}.docx'
    file_path = os.path.join(out_dir, file_name)

    doc = Document()

    # ---- 页面设置 (A4) ----
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ---- 封面标题 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('活动台账 会议录音总结')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    # 副标题
    subtitle_text = activity.content or '活动台账'
    if len(subtitle_text) > 60:
        subtitle_text = subtitle_text[:60] + '...'
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(subtitle_text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x60, 0x62, 0x66)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    activity_date = activity.date.strftime('%Y年%m月%d日') if activity.date else '未知日期'
    run = info.add_run(f'活动时间: {activity_date}    生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x90, 0x93, 0x99)

    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('━' * 50)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xc8, 0xda, 0xf0)
    doc.add_paragraph()

    # ---- 解析 Markdown → docx 的辅助 ----
    def _add_markdown_content(text, doc):
        """将 Markdown 文本解析后写入 Document"""
        in_table = False
        table = None
        lines = text.strip().split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                if in_table:
                    in_table = False
                    table = None
                doc.add_paragraph()
                continue

            # 分隔线
            if re.match(r'^[━═]{5,}', stripped):
                continue

            # 表格行
            if stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 2:
                if not re.match(r'^\|[-: ]+\|', stripped):
                    cells = [c.strip() for c in stripped.split('|')[1:-1]]
                    if len(cells) >= 2:
                        if not in_table:
                            table = doc.add_table(rows=0, cols=len(cells))
                            table.style = 'Table Grid'
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            in_table = True
                        row = table.add_row()
                        for i, cell_text in enumerate(cells):
                            row.cells[i].text = _strip_markdown(cell_text)
                            for paragraph in row.cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                        continue

            in_table = False
            table = None

            # 一级标题 一、 / # /
            if re.match(r'^[一二三四五六七八九十]、', stripped) or (stripped.startswith('# ') and not stripped.startswith('## ')):
                text_clean = re.sub(r'^#\s*', '', stripped)
                text_clean = _strip_markdown(text_clean)
                p = doc.add_paragraph()
                run = p.add_run(text_clean)
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = 'SimHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                continue

            # 二级标题 ## / 1. ** / 1、
            if stripped.startswith('## ') or re.match(r'^\d+[\.\、]\s*\*\*', stripped):
                text_clean = re.sub(r'^##\s*', '', stripped)
                text_clean = _strip_markdown(text_clean)
                p = doc.add_paragraph()
                run = p.add_run(text_clean)
                run.bold = True
                run.font.size = Pt(13)
                run.font.name = 'SimSun'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.color.rgb = RGBColor(0x30, 0x31, 0x33)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                continue

            # 列表项
            if re.match(r'^[-*]\s', stripped) or re.match(r'^\d+[\.\、]\s', stripped):
                text_clean = re.sub(r'^[-*]\s*|^\d+[\.\、]\s*', '', stripped)
                text_clean = _strip_markdown(text_clean)
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(text_clean)
                run.font.size = Pt(11)
                run.font.name = 'SimSun'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                continue

            # 普通段落
            text_clean = _strip_markdown(stripped)
            p = doc.add_paragraph()
            run = p.add_run(text_clean)
            run.font.size = Pt(12)
            run.font.name = 'SimSun'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5

    # ---- 第一部分: 发言分段 ----
    if segmented_text and segmented_text.strip():
        h1 = doc.add_paragraph()
        run = h1.add_run('一、发言分段原文')
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(8)
        _add_markdown_content(segmented_text, doc)

    # ---- 第二部分: 清洁版 ----
    if clean_text and clean_text.strip():
        h1 = doc.add_paragraph()
        run = h1.add_run('二、清洁版')
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(8)
        _add_markdown_content(clean_text, doc)

    # ---- 第三部分: 摘要版 ----
    if summary_text and summary_text.strip():
        h1 = doc.add_paragraph()
        run = h1.add_run('三、摘要版')
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(8)
        _add_markdown_content(summary_text, doc)

    # ---- 页脚 ----
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('—— 本报告由 AI 自动生成，仅供参考 ——')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x90, 0x93, 0x99)
    run.italic = True

    doc.save(file_path)

    static_url = f'/static/meetings/{file_name}'
    return static_url, file_name
