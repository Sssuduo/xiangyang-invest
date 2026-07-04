import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _ensure_dir():
    """确保输出目录存在"""
    base = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'assessments')
    os.makedirs(base, exist_ok=True)
    return base


def _set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        from lxml import etree
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        from lxml import etree
        element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')


def _strip_markdown(text):
    """移除 markdown 标记，保留纯文本"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def generate_assessment_docx(lead_name, enterprise_name, assessment_text):
    """将 AI 研判结果生成格式化的 Word 文档

    Args:
        lead_name: 项目名称
        enterprise_name: 企业名称
        assessment_text: LLM 返回的原始文本

    Returns:
        tuple: (file_path, file_name)
    """
    out_dir = _ensure_dir()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', lead_name)[:40]
    file_name = f'{safe_name}_AI研判报告_{timestamp}.docx'
    file_path = os.path.join(out_dir, file_name)

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ---- 设置默认字体 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ---- 标题 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('招商线索研判分析报告')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    # ---- 副标题 ----
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'项目：{lead_name}')
    run.font.size = Pt(14)
    run.font.name = 'SimSun'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(0x60, 0x62, 0x66)

    # ---- 生成信息 ----
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}    企业：{enterprise_name}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x90, 0x93, 0x99)

    doc.add_paragraph()  # 空行

    # ---- 分隔线 ----
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('━' * 50)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xc8, 0xda, 0xf0)

    doc.add_paragraph()

    # ---- 解析并生成内容 ----
    lines = assessment_text.strip().split('\n')
    in_table = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            in_table = False
            continue

        # 分隔线 → 一级标题
        if re.match(r'^[━═]{5,}', stripped):
            # 跳过分隔线，接下来的一行作为标题处理
            continue

        # 检测一级标题：一、 二、 或 1. 开头且长度较短
        is_section_title = re.match(r'^[一二三四五六七八九十]、', stripped)
        is_numbered_title = re.match(r'^\d+[\.\、]\s*\*\*', stripped)
        is_markdown_h1 = stripped.startswith('# ') and not stripped.startswith('## ')
        is_markdown_h2 = stripped.startswith('## ')

        if is_section_title or is_markdown_h1:
            text = re.sub(r'^#\s*', '', stripped)
            text = _strip_markdown(text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'SimHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            in_table = False
            continue

        if is_markdown_h2 or (is_numbered_title and len(stripped) < 80):
            text = re.sub(r'^##\s*', '', stripped)
            text = _strip_markdown(text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = 'SimSun'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.color.rgb = RGBColor(0x30, 0x31, 0x33)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            in_table = False
            continue

        # 检测表格行：| 列1 | 列2 | ... |
        if stripped.startswith('|') and stripped.endswith('|') and not re.match(r'^\|[-: ]+\|', stripped):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table and len(cells) >= 2:
                # 创建新表格
                table = doc.add_table(rows=0, cols=len(cells))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                in_table = True
            row = table.add_row()
            for i, cell_text in enumerate(cells):
                row.cells[i].text = _strip_markdown(cell_text)
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            continue

        # 普通段落
        in_table = False
        text = _strip_markdown(stripped)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'SimSun'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
        p.paragraph_format.line_spacing = 1.5

    # ---- 页脚 ----
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('—— 本报告由 AI 自动生成，仅供参考 ——')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x90, 0x93, 0x99)
    run.italic = True

    doc.save(file_path)

    # 返回相对于 static 目录的路径
    relative_dir = os.path.join('assessments', file_name)
    static_url = f'/static/assessments/{file_name}'
    return static_url, file_name
